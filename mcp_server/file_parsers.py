"""
File parser module with unified interface for multiple formats
"""
import os
import pandas as pd
import fitz  # PyMuPDF
from docx import Document
from pathlib import Path
from typing import Dict, Any, List, Optional
import json


class FileParserBase:
    def parse(self, file_path: str, **kwargs) -> Dict[str, Any]:
        raise NotImplementedError


class CSVParser(FileParserBase):
    
    def parse(self, file_path: str, **kwargs) -> Dict[str, Any]:
        max_rows = kwargs.get('max_rows', 1000)
        encoding = kwargs.get('encoding', 'utf-8')
        
        try:
            df = pd.read_csv(file_path, encoding=encoding)
            
            # Limit returned rows
            preview_df = df.head(max_rows) if len(df) > max_rows else df
            
            return {
                "type": "csv",
                "success": True,
                "file_info": {
                    "total_rows": len(df),
                    "total_columns": len(df.columns),
                    "file_size": os.path.getsize(file_path),
                    "columns": df.columns.tolist()
                },
                "data": {
                    "preview_rows": len(preview_df),
                    "records": preview_df.to_dict('records'),
                    "summary": {
                        "numeric_columns": df.select_dtypes(include=['number']).columns.tolist(),
                        "text_columns": df.select_dtypes(include=['object']).columns.tolist(),
                        "missing_values": df.isnull().sum().to_dict()
                    }
                }
            }
        except Exception as e:
            return {
                "type": "csv", 
                "success": False, 
                "error": str(e)
            }


class PDFParser(FileParserBase):
    
    def parse(self, file_path: str, **kwargs) -> Dict[str, Any]:
        max_pages = kwargs.get('max_pages', 10)
        
        try:
            doc = fitz.open(file_path)
            text_content = []
            
            for page_num in range(min(len(doc), max_pages)):
                page = doc[page_num]
                text = page.get_text()
                if text.strip():
                    text_content.append({
                        "page": page_num + 1,
                        "text": text.strip()
                    })
            
            doc.close()
            
            # Extract document metadata
            doc_info = fitz.open(file_path)
            metadata = doc_info.metadata
            doc_info.close()
            
            return {
                "type": "pdf",
                "success": True,
                "file_info": {
                    "total_pages": len(fitz.open(file_path)),
                    "file_size": os.path.getsize(file_path),
                    "metadata": {
                        "title": metadata.get('title', ''),
                        "author": metadata.get('author', ''),
                        "subject": metadata.get('subject', ''),
                        "creator": metadata.get('creator', '')
                    }
                },
                "data": {
                    "extracted_pages": len(text_content),
                    "content": text_content,
                    "word_count": sum(len(page['text'].split()) for page in text_content)
                }
            }
        except Exception as e:
            return {
                "type": "pdf",
                "success": False,
                "error": str(e)
            }


class DOCXParser(FileParserBase):
    
    def parse(self, file_path: str, **kwargs) -> Dict[str, Any]:
        try:
            doc = Document(file_path)
            
            # Extract paragraph text
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            # Extract table data
            tables_data = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    tables_data.append(table_data)
            
            return {
                "type": "docx",
                "success": True,
                "file_info": {
                    "file_size": os.path.getsize(file_path),
                    "paragraph_count": len(paragraphs),
                    "table_count": len(tables_data)
                },
                "data": {
                    "paragraphs": paragraphs,
                    "tables": tables_data,
                    "word_count": sum(len(p.split()) for p in paragraphs)
                }
            }
        except Exception as e:
            return {
                "type": "docx",
                "success": False,
                "error": str(e)
            }


class TXTParser(FileParserBase):
    
    def parse(self, file_path: str, **kwargs) -> Dict[str, Any]:
        encoding = kwargs.get('encoding', 'utf-8')
        max_chars = kwargs.get('max_chars', 10000)
        
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                content = f.read()
            
            # Limit returned characters
            preview_content = content[:max_chars] if len(content) > max_chars else content
            
            return {
                "type": "txt",
                "success": True,
                "file_info": {
                    "file_size": os.path.getsize(file_path),
                    "total_chars": len(content),
                    "total_lines": content.count('\n') + 1
                },
                "data": {
                    "content": preview_content,
                    "word_count": len(content.split()),
                    "line_count": preview_content.count('\n') + 1,
                    "is_truncated": len(content) > max_chars
                }
            }
        except Exception as e:
            return {
                "type": "txt",
                "success": False,
                "error": str(e)
            }


class FileParserManager:
    
    def __init__(self):
        self.parsers = {
            '.csv': CSVParser(),
            '.pdf': PDFParser(),
            '.docx': DOCXParser(),
            '.txt': TXTParser(),
            '.md': TXTParser(),  # Markdown files use txt parser
        }
    
    def get_supported_extensions(self) -> List[str]:
        return list(self.parsers.keys())
    
    def is_supported(self, file_path: str) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext in self.parsers
    
    def parse_file(self, file_path: str, **kwargs) -> Dict[str, Any]:
        if not os.path.exists(file_path):
            return {
                "success": False,
                "error": f"File not found: {file_path}"
            }
        
        ext = Path(file_path).suffix.lower()
        if ext not in self.parsers:
            return {
                "success": False,
                "error": f"Unsupported file type: {ext}",
                "supported_types": self.get_supported_extensions()
            }
        
        parser = self.parsers[ext]
        result = parser.parse(file_path, **kwargs)
        
        # Add common file info
        result["file_path"] = file_path
        result["file_name"] = os.path.basename(file_path)
        result["file_extension"] = ext
        
        return result
    
    def register_parser(self, extension: str, parser: FileParserBase):
        self.parsers[extension.lower()] = parser