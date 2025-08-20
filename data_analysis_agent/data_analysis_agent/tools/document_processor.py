"""
文档处理工具 - 支持多种文件格式的读取和预处理
"""
import pandas as pd
import numpy as np
from typing import Dict, Any, Optional, Tuple, List
import PyPDF2
import pdfplumber
from docx import Document
import json
import os
import io


class DocumentProcessor:
    """文档处理工具类"""
    
    def __init__(self):
        self.supported_formats = ['.csv', '.xlsx', '.xls', '.pdf', '.txt', '.json', '.docx']
    
    def read_file(self, file_path: str) -> Tuple[Any, Dict[str, Any]]:
        """
        读取文件并返回数据和元信息
        
        Args:
            file_path: 文件路径
            
        Returns:
            (data, info): 数据内容和文件信息
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        file_size = os.path.getsize(file_path)
        
        info = {
            'file_path': file_path,
            'file_type': file_ext,
            'file_size': file_size,
            'file_name': os.path.basename(file_path)
        }
        
        if file_ext == '.csv':
            return self._read_csv(file_path, info)
        elif file_ext in ['.xlsx', '.xls']:
            return self._read_excel(file_path, info)
        elif file_ext == '.pdf':
            return self._read_pdf(file_path, info)
        elif file_ext == '.txt':
            return self._read_txt(file_path, info)
        elif file_ext == '.json':
            return self._read_json(file_path, info)
        elif file_ext == '.docx':
            return self._read_docx(file_path, info)
        else:
            raise ValueError(f"不支持的文件格式: {file_ext}")
    
    def _read_csv(self, file_path: str, info: Dict[str, Any]) -> Tuple[pd.DataFrame, Dict[str, Any]]:
        """读取CSV文件"""
        try:
            # 尝试不同的编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin1']
            df = None
            
            for encoding in encodings:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    info['encoding'] = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if df is None:
                raise ValueError("无法解析CSV文件编码")
            
            # 添加数据信息
            info.update({
                'shape': df.shape,
                'columns': df.columns.tolist(),
                'dtypes': df.dtypes.to_dict(),
                'missing_values': df.isnull().sum().to_dict(),
                'memory_usage': df.memory_usage().sum()
            })
            
            return df, info
            
        except Exception as e:
            raise ValueError(f"读取CSV文件失败: {str(e)}")
    
    def _read_excel(self, file_path: str, info: Dict[str, Any]) -> Tuple[pd.DataFrame, Dict[str, Any]]:
        """读取Excel文件"""
        try:
            # 读取第一个sheet
            df = pd.read_excel(file_path)
            
            # 获取所有sheet名称
            with pd.ExcelFile(file_path) as xls:
                sheet_names = xls.sheet_names
            
            info.update({
                'sheet_names': sheet_names,
                'shape': df.shape,
                'columns': df.columns.tolist(),
                'dtypes': df.dtypes.to_dict(),
                'missing_values': df.isnull().sum().to_dict()
            })
            
            return df, info
            
        except Exception as e:
            raise ValueError(f"读取Excel文件失败: {str(e)}")
    
    def _read_pdf(self, file_path: str, info: Dict[str, Any]) -> Tuple[str, Dict[str, Any]]:
        """读取PDF文件"""
        try:
            text_content = ""
            
            # 使用pdfplumber提取文本
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    text_content += page.extract_text() or ""
                    text_content += "\n"
            
            info.update({
                'page_count': page_count,
                'text_length': len(text_content),
                'word_count': len(text_content.split())
            })
            
            return text_content, info
            
        except Exception as e:
            # 如果pdfplumber失败，尝试PyPDF2
            try:
                text_content = ""
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    page_count = len(pdf_reader.pages)
                    
                    for page in pdf_reader.pages:
                        text_content += page.extract_text()
                        text_content += "\n"
                
                info.update({
                    'page_count': page_count,
                    'text_length': len(text_content),
                    'word_count': len(text_content.split())
                })
                
                return text_content, info
            except Exception as e2:
                raise ValueError(f"读取PDF文件失败: {str(e2)}")
    
    def _read_txt(self, file_path: str, info: Dict[str, Any]) -> Tuple[str, Dict[str, Any]]:
        """读取文本文件"""
        try:
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin1']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    info['encoding'] = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise ValueError("无法解析文本文件编码")
            
            info.update({
                'text_length': len(content),
                'line_count': content.count('\n') + 1,
                'word_count': len(content.split())
            })
            
            return content, info
            
        except Exception as e:
            raise ValueError(f"读取文本文件失败: {str(e)}")
    
    def _read_json(self, file_path: str, info: Dict[str, Any]) -> Tuple[Any, Dict[str, Any]]:
        """读取JSON文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            info.update({
                'data_type': type(data).__name__,
                'keys': list(data.keys()) if isinstance(data, dict) else None,
                'length': len(data) if hasattr(data, '__len__') else None
            })
            
            return data, info
            
        except Exception as e:
            raise ValueError(f"读取JSON文件失败: {str(e)}")
    
    def _read_docx(self, file_path: str, info: Dict[str, Any]) -> Tuple[str, Dict[str, Any]]:
        """读取Word文档"""
        try:
            doc = Document(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            info.update({
                'paragraph_count': len(doc.paragraphs),
                'text_length': len(text_content),
                'word_count': len(text_content.split())
            })
            
            return text_content, info
            
        except Exception as e:
            raise ValueError(f"读取Word文档失败: {str(e)}")
    
    def get_data_summary(self, data: Any, info: Dict[str, Any]) -> str:
        """
        生成数据摘要描述
        
        Args:
            data: 数据内容
            info: 文件信息
            
        Returns:
            数据摘要文本
        """
        summary_parts = []
        
        summary_parts.append(f"文件名: {info['file_name']}")
        summary_parts.append(f"文件类型: {info['file_type']}")
        summary_parts.append(f"文件大小: {self._format_size(info['file_size'])}")
        
        if isinstance(data, pd.DataFrame):
            summary_parts.append(f"数据形状: {info['shape'][0]}行 × {info['shape'][1]}列")
            summary_parts.append(f"列名: {', '.join(info['columns'][:5])}{'...' if len(info['columns']) > 5 else ''}")
            
            # 缺失值统计
            missing_total = sum(info['missing_values'].values())
            if missing_total > 0:
                summary_parts.append(f"缺失值总数: {missing_total}")
        
        elif isinstance(data, str):
            if 'word_count' in info:
                summary_parts.append(f"文本长度: {info['word_count']}词")
            if 'page_count' in info:
                summary_parts.append(f"页数: {info['page_count']}")
        
        return "\n".join(summary_parts)
    
    def _format_size(self, size_bytes: int) -> str:
        """格式化文件大小"""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f} TB"